# parse_docx.py —— Word（.docx）解析：pandoc 主解析 + python-docx 降级。
#
# 能力：
#   - pandoc docx→markdown：保留标题层级；OMML 公式自动转 $LaTeX$（教育场景刚需）
#   - --extract-media 导出内嵌图片到媒体目录，markdown 改写为持久引用
#   - pandoc 缺失/失败时降级 python-docx（纯文本，公式丢失并警告）
#
# 注意：老格式 .doc（OLE 二进制）pandoc 不支持，由 rag 侧明确拒绝并提示另存 .docx。
#
# usage: python3 parse_docx.py <input_abs> <out_json_abs> <media_dir_abs>
import glob
import os
import re
import shutil
import subprocess
import sys
import tempfile

import common


def _pandoc(input_path: str, tmp: str) -> str:
    """调 pandoc 转 markdown，返回 md 文本。pandoc 不可用/失败抛异常。"""
    out_md = os.path.join(tmp, "out.md")
    cmd = [
        "pandoc", input_path,
        "-t", "markdown",
        "--extract-media=" + os.path.join(tmp, "media"),
        "-o", out_md,
    ]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=110)
    if proc.returncode != 0:
        raise RuntimeError("pandoc 失败: %s" % (proc.stderr or proc.stdout)[:500])
    with open(out_md, encoding="utf-8") as f:
        return f.read()


def _collect_media(tmp_media: str, media_dir: str, media_ref: str) -> list[dict]:
    """把 pandoc 导出的图片移动到持久媒体目录，返回 media 清单。"""
    if not os.path.isdir(tmp_media):
        return []
    common.ensure_media_dir(media_dir)
    media = []
    for f in glob.glob(os.path.join(tmp_media, "**", "*"), recursive=True):
        if not os.path.isfile(f):
            continue
        name = common.safe_filename(os.path.basename(f))
        base, ext = os.path.splitext(name)
        dst = os.path.join(media_dir, name)
        i = 1
        while os.path.exists(dst):
            dst = os.path.join(media_dir, "%s_%d%s" % (base, i, ext))
            i += 1
        shutil.move(f, dst)
        ref = os.path.join(media_ref, os.path.basename(dst)).replace(os.sep, "/")
        media.append({"type": common.media_type(name), "path": ref, "alt": ""})
    return media


def _fallback_python_docx(input_path: str) -> str:
    """pandoc 不可用时的降级：python-docx 提取段落文本（公式丢失）。"""
    from docx import Document  # python-docx
    doc = Document(input_path)
    md = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            md.append("#" * min(level, 6) + " " + text)
        else:
            md.append(text)
    # 表格 → markdown 表格
    for tbl in doc.tables:
        md.append("")
        for row in tbl.rows:
            cells = [c.text.strip().replace("|", "\\|") for c in row.cells]
            md.append("| " + " | ".join(cells) + " |")
        md.append("")
    return "\n\n".join(md)


def main() -> int:
    if len(sys.argv) != 4:
        common.usage("parse_docx.py")
    input_path, out_path, media_dir = sys.argv[1], sys.argv[2], sys.argv[3]
    media_ref = common.media_rel_prefix(media_dir)

    warnings: list[str] = []
    media: list[dict] = []
    tmp = tempfile.mkdtemp(prefix="docx_")
    try:
        try:
            md = _pandoc(input_path, tmp)
        except FileNotFoundError:
            warnings.append("沙盒缺少 pandoc，已降级纯文本提取，公式（LaTeX）丢失")
            md = _fallback_python_docx(input_path)
        except Exception as e:  # noqa: BLE001
            warnings.append("pandoc 解析失败（%s），已降级纯文本提取" % str(e)[:120])
            try:
                md = _fallback_python_docx(input_path)
            except Exception as e2:  # noqa: BLE001
                common.emit(out_path, "", "", [], False, ["docx 解析失败: %s" % e2])
                return 0

        # 图片导出 & 媒体引用改写（pandoc 成功且确有导出时）。
        tmp_media = os.path.join(tmp, "media")
        if os.path.isdir(tmp_media):
            media = _collect_media(tmp_media, media_dir, media_ref)
            md = common.rewrite_media_refs(md, media_ref)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    # 标题：取首个一级标题。
    title = ""
    for line in md.splitlines():
        m = re.match(r"^#\s+(.+)$", line)
        if m:
            title = m.group(1).strip()
            break

    common.emit(out_path, title, md, media, False, warnings)
    return 0


if __name__ == "__main__":
    sys.exit(main())
