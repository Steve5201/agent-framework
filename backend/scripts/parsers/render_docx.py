# render_docx.py —— 按 DocumentSpec 渲染 Word（.docx）文档（P4-D / P4-J 阶段2）。
#
# 技术栈：python-docx（沙盒镜像已预装 python-docx==1.1.2，见 deploy/Dockerfile）。
# 能力：
#   - 封面标题/副标题居中，正文 Heading 1 章节 + 富文本块渲染（P4-J 阶段2）：
#     blocks 六种类型——paragraph / list / table / image / formula / code；
#   - blocks 为空时回退旧契约 body 轻量标记（# 二级 / ## 三级 / - 列表 / 段落）；
#   - 行内 **加粗** 转真实加粗 run（python-docx 无 markdown 支持，自解析）；
#   - image 块支持知识库媒体（/work/rag-media/…，Go 侧解析为绝对路径）与
#     内联 SVG（svg://…，PyMuPDF 转 PNG）；formula 块用 matplotlib mathtext
#     渲染为透明 PNG 后插入（沙盒需预装 matplotlib，见 deploy/Dockerfile）；
#   - 页脚文本（文档末尾居中）。
#
# usage: python3 render_docx.py <spec_json_abs> <out_docx_abs>
import os
import re
import sys

import render_common


def _add_bold_runs(paragraph, text: str) -> None:
    """把含 **加粗** 标记的文本拆成普通/加粗 run 写入段落。"""
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, seg in enumerate(parts):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        if i % 2 == 1:
            run.bold = True


def _set_east_asia(paragraph, name: str = "微软雅黑", latin: str = None) -> None:
    """设置中文字体（w:eastAsia）+ 可选拉丁字体：字体在客户端打开时解析，缺省会乱码。"""
    from docx.oxml.ns import qn

    for run in paragraph.runs:
        run.font.name = latin or name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), name)


def _render_table(doc, b: dict) -> None:
    """表格块：首行表头加粗 + Table Grid 边框。"""
    from docx.shared import Pt

    headers = b.get("headers") or []
    rows = b.get("rows") or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
            _set_east_asia(p)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
            for p in cell.paragraphs:
                _set_east_asia(p)


def _render_image(doc, b: dict, work_dir: str) -> None:
    """图片块：知识库媒体 / 内联 SVG → PNG 插入（居中），附小字图注。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    try:
        png, _, _ = render_common.asset_png(b.get("src", ""), work_dir)
    except ValueError as e:
        p = doc.add_paragraph("（图片无法加载：%s）" % e)
        _set_east_asia(p)
        return
    width_in = Inches(5) if b.get("width", 0) <= 0 else Inches(b["width"] / 96.0)
    doc.add_picture(png, width=width_in)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if b.get("caption"):
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(b["caption"])
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        _set_east_asia(cap)


def _render_formula(doc, b: dict, work_dir: str) -> None:
    """公式块：render=image → matplotlib PNG；render=native → OMML 原生公式
    （可编辑/可复制，LaTeX 常见子集），结构不支持时自动回退图片渲染。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    if b.get("render", "image") == "native":
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            render_common.omml_into_docx_paragraph(
                p, render_common.latex_to_omml(b.get("text", ""))
            )
            return
        except Exception:
            pass  # OMML 注入失败 → 回退图片
    try:
        png = render_common.formula_png(b.get("text", ""), work_dir)
    except ValueError as e:
        p = doc.add_paragraph("（公式渲染失败：%s）" % e)
        _set_east_asia(p)
        return
    doc.add_picture(png, width=Inches(3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _render_code(doc, b: dict) -> None:
    """代码块：等宽字体逐行段落（保留缩进）。"""
    from docx.shared import Pt

    for line in (b.get("text", "") or "").splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        _set_east_asia(p, "微软雅黑", latin="Consolas")


def _render_blocks(doc, blocks, work_dir: str) -> None:
    """按 blocks 顺序渲染章节内容（P4-J 阶段2 六种块 + 旧契约 kind 兼容）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for b in blocks:
        kind = b.get("type", "")
        if kind in ("paragraph", "p"):
            p = doc.add_paragraph()
            _add_bold_runs(p, b.get("text", ""))
            _set_east_asia(p)
        elif kind == "h2":
            p = doc.add_heading(b.get("text", ""), level=2)
            _set_east_asia(p)
        elif kind == "h3":
            p = doc.add_heading(b.get("text", ""), level=3)
            _set_east_asia(p)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_bold_runs(p, b.get("text", ""))
            _set_east_asia(p)
        elif kind == "list":
            for item in b.get("items") or []:
                p = doc.add_paragraph(style="List Bullet")
                _add_bold_runs(p, item)
                _set_east_asia(p)
        elif kind == "table":
            _render_table(doc, b)
        elif kind == "image":
            _render_image(doc, b, work_dir)
        elif kind == "formula":
            _render_formula(doc, b, work_dir)
        elif kind == "code":
            _render_code(doc, b)
        else:
            p = doc.add_paragraph("（忽略不支持的块类型：%s）" % kind)
            _set_east_asia(p)


def main() -> int:
    if len(sys.argv) != 3:
        render_common.usage("render_docx.py")
    spec_path, out_path = sys.argv[1], sys.argv[2]
    try:
        spec = render_common.load_spec(spec_path)
    except ValueError as e:
        render_common.emit_error("render_docx: %s" % e)
        return 1

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        render_common.emit_error("render_docx: 沙盒缺少 python-docx（%s）" % e)
        return 1

    # 工作目录（脚本 cwd = /work/users/<uid>）：asset/formula 临时 PNG 落此。
    work_dir = os.getcwd()
    doc = Document()

    # 封面标题（居中，大字号）。
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(spec["title"])
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    _set_east_asia(title_p)

    if spec["subtitle"]:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(spec["subtitle"])
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        _set_east_asia(sub_p)

    # 章节正文：blocks 非空优先，否则回退 body 轻量标记。
    for sec in spec["sections"]:
        h = doc.add_heading(sec["heading"], level=1)
        _set_east_asia(h)
        _render_blocks(doc, render_common.blocks_of(sec), work_dir)

    if spec["footer"]:
        foot_p = doc.add_paragraph()
        foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        foot_run = foot_p.add_run(spec["footer"])
        foot_run.font.size = Pt(10)
        foot_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        _set_east_asia(foot_p)

    try:
        doc.save(out_path)
    except OSError as e:
        render_common.emit_error("render_docx: 保存失败（%s）" % e)
        return 1

    size = render_common.file_size(out_path)
    if size <= 0:
        render_common.emit_error("render_docx: 产物为空或不可读")
        return 1
    render_common.emit_ok(out_path, size)
    return 0


if __name__ == "__main__":
    sys.exit(main())
